from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.common import DocumentChunk, KnowledgeDocument
from app.services.embeddings import EmbeddingService


@dataclass
class ParsedDocument:
    title: str
    text: str
    file_type: str


class DocumentIngestionService:
    supported_suffixes = {".txt", ".md", ".docx", ".pdf"}

    def __init__(self, db: AsyncSession, embedding_service: EmbeddingService | None = None) -> None:
        self.db = db
        self.embedding_service = embedding_service or EmbeddingService()

    @staticmethod
    def default_knowledge_base_dir() -> Path:
        return Path(__file__).resolve().parents[3] / "Knowledge Base"

    @classmethod
    def knowledge_base_files(cls, directory: Path | None = None) -> list[Path]:
        knowledge_base_dir = directory or cls.default_knowledge_base_dir()
        if not knowledge_base_dir.exists():
            raise FileNotFoundError(f"Knowledge Base directory not found: {knowledge_base_dir}")
        return sorted(
            path
            for path in knowledge_base_dir.iterdir()
            if path.is_file()
            and not path.name.startswith((".", "~$"))
            and path.suffix.lower() in cls.supported_suffixes
        )

    async def sync_knowledge_base(self, directory: Path | None = None) -> list[KnowledgeDocument]:
        """Index local knowledge-base files that are not already in the database."""
        existing_paths = set(
            (await self.db.scalars(select(KnowledgeDocument.source_path))).all()
        )
        ingested: list[KnowledgeDocument] = []

        for path in self.knowledge_base_files(directory):
            if path.name in existing_paths:
                continue
            ingested.append(await self.ingest_bytes(path.read_bytes(), path.name, title=path.stem))

        return ingested

    def _chunk_text(self, text: str, max_chars: int = 1000, overlap: int = 120) -> list[str]:
        cleaned = "\n".join(line.strip() for line in text.splitlines()).strip()
        if not cleaned:
            return []
        chunks: list[str] = []
        start = 0
        while start < len(cleaned):
            end = min(len(cleaned), start + max_chars)
            chunks.append(cleaned[start:end].strip())
            if end >= len(cleaned):
                break
            start = max(end - overlap, start + 1)
        return [chunk for chunk in chunks if chunk]

    def _parse_txt(self, content: bytes, filename: str) -> ParsedDocument:
        return ParsedDocument(title=Path(filename).stem, text=content.decode("utf-8", errors="ignore"), file_type="txt")

    def _parse_docx(self, content: bytes, filename: str) -> ParsedDocument:
        doc = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return ParsedDocument(title=Path(filename).stem, text=text, file_type="docx")

    def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        reader = PdfReader(BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return ParsedDocument(title=Path(filename).stem, text=text, file_type="pdf")

    def parse_document(self, content: bytes, filename: str) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix in {".txt", ".md"}:
            return self._parse_txt(content, filename)
        if suffix == ".docx":
            return self._parse_docx(content, filename)
        if suffix == ".pdf":
            return self._parse_pdf(content, filename)
        raise ValueError(f"Unsupported file type: {suffix}")

    async def ingest_bytes(self, content: bytes, filename: str, title: str | None = None) -> KnowledgeDocument:
        parsed = self.parse_document(content, filename)
        chunks = self._chunk_text(parsed.text)
        document = KnowledgeDocument(title=title or parsed.title, source_path=filename, file_type=parsed.file_type)
        self.db.add(document)
        await self.db.flush()

        for index, chunk in enumerate(chunks):
            embedding = self.embedding_service.embed_query(chunk)
            self.db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=index,
                    content=chunk,
                    page_number=None,
                    chunk_metadata={"title": document.title, "source_path": filename, "file_type": parsed.file_type},
                    embedding=embedding,
                )
            )

        await self.db.flush()
        return document
